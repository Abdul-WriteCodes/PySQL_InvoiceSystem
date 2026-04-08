import streamlit as st
import pandas as pd
import numpy as np
import plotly.express as px
import plotly.graph_objects as go
from plotly.subplots import make_subplots
import requests
import json
import io
import warnings
warnings.filterwarnings("ignore")


# ── Google Sheets credit system ───────────────────────────────────────────────
from google.oauth2.service_account import Credentials
import gspread

GSHEET_SCOPES = [
    "https://www.googleapis.com/auth/spreadsheets",
    "https://www.googleapis.com/auth/drive",
]

def _get_sheet():
    """Authenticate and return the credits worksheet."""
    creds_dict = dict(st.secrets["gcp_service_account"])
    creds = Credentials.from_service_account_info(creds_dict, scopes=GSHEET_SCOPES)
    gc = gspread.authorize(creds)
    spreadsheet_id = st.secrets["SHEET_ID"]
    sh = gc.open_by_key(spreadsheet_id)
    return sh.sheet1   # first sheet: Key | Credits | DatePurchased | Email


def lookup_key(access_key: str) -> dict | None:
    """
    Find a row by access key.
    Returns dict with keys: row_index, key, credits, date_purchased, email
    or None if not found.
    """
    try:
        ws = _get_sheet()
        records = ws.get_all_records()           # [{Key, Credits, DatePurchased, Email}, …]
        for i, row in enumerate(records, start=2):  # row 1 is header
            if str(row.get("Key", "")).strip() == access_key.strip():
                return {
                    "row_index": i,
                    "key": row["Key"],
                    "credits": int(row.get("Credits", 0)),
                    "date_purchased": row.get("DatePurchased", ""),
                    "email": row.get("Email", ""),
                }
        return None
    except Exception as e:
        st.error(f"Sheet lookup error: {e}")
        return None


def deduct_credit(row_index: int, current_credits: int) -> int:
    """Write credits − 1 back to the sheet. Returns new credit count."""
    try:
        ws = _get_sheet()
        new_credits = max(0, current_credits - 1)
        # Column B = Credits (column index 2)
        ws.update_cell(row_index, 2, new_credits)
        return new_credits
    except Exception as e:
        st.error(f"Credit deduction error: {e}")
        return current_credits

# ── Page config ──────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="PanelStatX",
    page_icon="⬡",
    layout="wide",
    initial_sidebar_state="expanded",
)

# ── Custom CSS ────────────────────────────────────────────────────────────────
st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=DM+Mono:ital,wght@0,300;0,400;0,500;1,300&family=Syne:wght@400;500;600;700;800&display=swap');

:root {
    --bg:        #0a0c10;
    --surface:   #111318;
    --surface2:  #181c24;
    --border:    #1f2535;
    --accent:    #00e5c8;
    --accent2:   #7c6df0;
    --accent3:   #f05c7c;
    --text:      #e2e8f4;
    --muted:     #6b7a9a;
    --success:   #22d3a0;
    --warn:      #f5a623;
}

html, body, [data-testid="stAppViewContainer"] {
    background-color: var(--bg) !important;
    color: var(--text) !important;
    font-family: 'DM Mono', monospace !important;
}

[data-testid="stSidebar"] {
    background: var(--surface) !important;
    border-right: 1px solid var(--border) !important;
}

/* Headers */
h1, h2, h3, h4 {
    font-family: 'Syne', sans-serif !important;
    color: var(--text) !important;
}

/* Metric cards */
[data-testid="metric-container"] {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
    padding: 16px !important;
}
[data-testid="metric-container"] > div > div:first-child {
    color: var(--muted) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.72rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.1em !important;
}
[data-testid="metric-container"] label {
    color: var(--muted) !important;
    font-size: 0.72rem !important;
    text-transform: uppercase !important;
    letter-spacing: 0.1em !important;
}
[data-testid="stMetricValue"] {
    color: var(--accent) !important;
    font-family: 'Syne', sans-serif !important;
    font-weight: 700 !important;
}

/* Buttons */
.stButton > button {
    background: transparent !important;
    border: 1px solid var(--accent) !important;
    color: var(--accent) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.8rem !important;
    border-radius: 4px !important;
    padding: 8px 20px !important;
    transition: all 0.2s !important;
    letter-spacing: 0.05em !important;
}
.stButton > button:hover {
    background: var(--accent) !important;
    color: var(--bg) !important;
}

/* Primary button */
[data-testid="baseButton-primary"] > button,
.stButton [kind="primary"] {
    background: var(--accent) !important;
    color: var(--bg) !important;
    font-weight: 600 !important;
}

/* Selectbox / inputs */
.stSelectbox > div > div,
.stMultiSelect > div > div,
.stTextInput > div > div > input,
.stNumberInput > div > div > input {
    background: var(--surface2) !important;
    border: 1px solid var(--border) !important;
    color: var(--text) !important;
    font-family: 'DM Mono', monospace !important;
    border-radius: 4px !important;
}

/* Tabs */
.stTabs [data-baseweb="tab-list"] {
    background: var(--surface) !important;
    border-bottom: 1px solid var(--border) !important;
    gap: 0 !important;
}
.stTabs [data-baseweb="tab"] {
    color: var(--muted) !important;
    font-family: 'DM Mono', monospace !important;
    font-size: 0.8rem !important;
    background: transparent !important;
    border-radius: 0 !important;
    padding: 12px 24px !important;
}
.stTabs [aria-selected="true"] {
    color: var(--accent) !important;
    border-bottom: 2px solid var(--accent) !important;
    background: transparent !important;
}

/* Dataframes */
.stDataFrame {
    border: 1px solid var(--border) !important;
    border-radius: 8px !important;
}

/* Expanders */
.streamlit-expanderHeader {
    background: var(--surface2) !important;
    color: var(--text) !important;
    font-family: 'DM Mono', monospace !important;
    border: 1px solid var(--border) !important;
    border-radius: 4px !important;
}

/* Sidebar labels */
.stSidebar label, .stSidebar .stMarkdown {
    color: var(--muted) !important;
    font-size: 0.78rem !important;
}

/* Slider */
.stSlider [data-baseweb="slider"] div[role="slider"] {
    background: var(--accent) !important;
}

/* Custom badge */
.badge {
    display: inline-block;
    padding: 2px 10px;
    border-radius: 20px;
    font-size: 0.7rem;
    font-family: 'DM Mono', monospace;
    letter-spacing: 0.05em;
    text-transform: uppercase;
}
.badge-teal  { background: rgba(0,229,200,0.12); color: var(--accent); border: 1px solid rgba(0,229,200,0.3); }
.badge-purple{ background: rgba(124,109,240,0.12); color: var(--accent2); border: 1px solid rgba(124,109,240,0.3); }
.badge-red   { background: rgba(240,92,124,0.12); color: var(--accent3); border: 1px solid rgba(240,92,124,0.3); }
.badge-warn  { background: rgba(245,166,35,0.12); color: var(--warn); border: 1px solid rgba(245,166,35,0.3); }

/* AI box */
.ai-box {
    background: linear-gradient(135deg, rgba(0,229,200,0.05) 0%, rgba(124,109,240,0.08) 100%);
    border: 1px solid rgba(0,229,200,0.25);
    border-left: 3px solid var(--accent);
    border-radius: 8px;
    padding: 20px 24px;
    font-family: 'DM Mono', monospace;
    font-size: 0.85rem;
    line-height: 1.75;
    color: var(--text);
    white-space: pre-wrap;
}
.ai-label {
    font-family: 'Syne', sans-serif;
    font-size: 0.65rem;
    letter-spacing: 0.15em;
    text-transform: uppercase;
    color: var(--accent);
    margin-bottom: 10px;
    display: flex;
    align-items: center;
    gap: 6px;
}
.spinner-dot::after { content: '●'; animation: blink 1s infinite; }
@keyframes blink { 0%,100%{opacity:1} 50%{opacity:0.2} }

/* Hero header */
.hero {
    padding: 28px 0 20px 0;
    border-bottom: 1px solid var(--border);
    margin-bottom: 28px;
}
.hero-title {
    font-family: 'Syne', sans-serif;
    font-size: 2.2rem;
    font-weight: 800;
    letter-spacing: -0.02em;
    color: var(--text);
    margin: 0;
    line-height: 1;
}
.hero-title span { color: var(--accent); }
.hero-sub {
    font-family: 'DM Mono', monospace;
    font-size: 0.78rem;
    color: var(--muted);
    margin-top: 6px;
    letter-spacing: 0.04em;
}
.stat-pill {
    display: inline-flex;
    align-items: center;
    gap: 6px;
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: 4px;
    padding: 4px 12px;
    font-size: 0.72rem;
    color: var(--muted);
    margin-right: 8px;
    margin-top: 12px;
    font-family: 'DM Mono', monospace;
}
.stat-pill b { color: var(--accent); }

/* Section card */
.scard {
    background: var(--surface2);
    border: 1px solid var(--border);
    border-radius: 8px;
    padding: 20px;
    margin-bottom: 16px;
}
.scard-title {
    font-family: 'Syne', sans-serif;
    font-size: 0.85rem;
    font-weight: 600;
    text-transform: uppercase;
    letter-spacing: 0.08em;
    color: var(--muted);
    margin-bottom: 14px;
}

/* Dividers */
hr { border-color: var(--border) !important; }

/* Info boxes */
[data-testid="stInfo"] { background: rgba(0,229,200,0.06) !important; border-left-color: var(--accent) !important; }
[data-testid="stWarning"] { background: rgba(245,166,35,0.06) !important; border-left-color: var(--warn) !important; }
[data-testid="stSuccess"] { background: rgba(34,211,160,0.06) !important; border-left-color: var(--success) !important; }

/* Scrollbar */
::-webkit-scrollbar { width: 4px; height: 4px; }
::-webkit-scrollbar-track { background: var(--bg); }
::-webkit-scrollbar-thumb { background: var(--border); border-radius: 2px; }
</style>
""", unsafe_allow_html=True)


# ═══════════════════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

PLOTLY_THEME = dict(
    paper_bgcolor="rgba(0,0,0,0)",
    plot_bgcolor="rgba(0,0,0,0)",
    font=dict(family="DM Mono, monospace", color="#6b7a9a", size=11),
    xaxis=dict(gridcolor="#1f2535", linecolor="#1f2535", zerolinecolor="#1f2535"),
    yaxis=dict(gridcolor="#1f2535", linecolor="#1f2535", zerolinecolor="#1f2535"),
    colorway=["#00e5c8", "#7c6df0", "#f05c7c", "#f5a623", "#22d3a0", "#60a5fa"],
    margin=dict(l=40, r=20, t=40, b=40),
)


def apply_theme(fig):
    fig.update_layout(**PLOTLY_THEME)
    return fig


def generate_demo_panel():
    """Generate a balanced panel dataset with realistic variation."""
    np.random.seed(42)
    n_entities, n_periods = 30, 10
    entities = [f"Entity_{i:02d}" for i in range(1, n_entities + 1)]
    years = list(range(2014, 2014 + n_periods))
    rows = []
    for e in entities:
        fe = np.random.randn()  # entity fixed effect
        for y in years:
            te = 0.05 * (y - 2014)  # time trend
            x1 = np.random.randn() + fe * 0.3
            x2 = np.random.uniform(0, 10)
            x3 = np.random.choice([0, 1], p=[0.6, 0.4])
            y_val = 2 + 0.8 * x1 - 0.4 * x2 + 1.2 * x3 + fe + te + np.random.randn() * 0.5
            rows.append({"entity": e, "year": y, "y": round(y_val, 4),
                         "x1": round(x1, 4), "x2": round(x2, 4), "x3": int(x3)})
    return pd.DataFrame(rows)


def run_ols(df, y_col, x_cols):
    from numpy.linalg import lstsq
    from scipy import stats as sc_stats
    X = np.column_stack([np.ones(len(df))] + [df[c].values for c in x_cols])
    y = df[y_col].values
    coeffs, residuals, rank, sv = lstsq(X, y, rcond=None)
    y_hat = X @ coeffs
    resid = y - y_hat
    n, k = X.shape
    k_vars = k - 1           # number of regressors (excl. intercept)
    dof = n - k
    s2 = np.sum(resid**2) / dof
    cov = s2 * np.linalg.inv(X.T @ X)
    se = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)
    ss_tot = np.sum((y - y.mean())**2)
    ss_res = np.sum(resid**2)
    ss_reg = ss_tot - ss_res
    r2 = 1 - ss_res / ss_tot
    r2_adj = 1 - (1 - r2) * (n - 1) / dof
    # F-statistic: (SS_reg / k_vars) / (SS_res / dof)
    f_stat = (ss_reg / max(k_vars, 1)) / (ss_res / max(dof, 1))
    f_p    = 1 - sc_stats.f.cdf(f_stat, dfn=k_vars, dfd=dof)
    names = ["const"] + list(x_cols)
    result_df = pd.DataFrame({"Variable": names, "Coeff": coeffs, "Std_Err": se,
                               "t_stat": t_stats, "p_value": p_vals})
    stats = {"R2": r2, "R2_adj": r2_adj, "N": n, "k": k_vars,
             "AIC": n * np.log(ss_res / n) + 2 * k,
             "BIC": n * np.log(ss_res / n) + k * np.log(n),
             "F_stat": f_stat, "F_p": f_p}
    return result_df, resid, y_hat, stats, cov


def run_within(df, y_col, x_cols, entity_col, time_col):
    """Fixed-effects (within) estimator via demeaning."""
    from scipy import stats as sc_stats
    panel = df.copy()
    for col in [y_col] + list(x_cols):
        entity_means = panel.groupby(entity_col)[col].transform("mean")
        time_means = panel.groupby(time_col)[col].transform("mean")
        grand_mean = panel[col].mean()
        panel[col + "_dm"] = panel[col] - entity_means - time_means + grand_mean
    y_dm = panel[y_col + "_dm"].values
    X_dm = np.column_stack([panel[c + "_dm"].values for c in x_cols])
    from numpy.linalg import lstsq
    coeffs, _, _, _ = lstsq(X_dm, y_dm, rcond=None)
    y_hat_dm = X_dm @ coeffs
    resid = y_dm - y_hat_dm
    n, k = X_dm.shape
    dof = n - k - df[entity_col].nunique() - df[time_col].nunique() + 1
    if dof <= 0:
        dof = max(1, n - k)
    s2 = np.sum(resid**2) / dof
    cov = s2 * np.linalg.inv(X_dm.T @ X_dm)
    se = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)
    ss_tot = np.sum((y_dm - y_dm.mean())**2)
    ss_res = np.sum(resid**2)
    ss_reg = ss_tot - ss_res
    r2 = max(0, 1 - ss_res / ss_tot)
    r2_adj = max(0, 1 - (1 - r2) * (n - 1) / dof)
    f_stat = (ss_reg / max(k, 1)) / (ss_res / max(dof, 1))
    f_p    = 1 - sc_stats.f.cdf(f_stat, dfn=k, dfd=dof)
    result_df = pd.DataFrame({"Variable": list(x_cols), "Coeff": coeffs,
                               "Std_Err": se, "t_stat": t_stats, "p_value": p_vals})
    stats = {"R2": r2, "R2_adj": r2_adj, "N": n, "k": k,
             "AIC": n * np.log(max(ss_res, 1e-10) / n) + 2 * k,
             "BIC": n * np.log(max(ss_res, 1e-10) / n) + k * np.log(n),
             "F_stat": f_stat, "F_p": f_p}
    return result_df, resid, y_hat_dm, stats, cov


def run_re(df, y_col, x_cols, entity_col, time_col):
    """
    Random Effects estimator (Swamy-Arora / GLS).
    Estimates the between-entity variance (sigma_u²) and within-entity
    variance (sigma_e²), computes the GLS theta weight, then runs quasi-
    demeaned OLS to obtain RE coefficients, SEs, t-stats, and p-values.
    Also returns the coefficient vcov for use in the Hausman test.
    """
    from numpy.linalg import lstsq
    from scipy import stats as sc_stats

    panel = df.copy().sort_values([entity_col, time_col])
    n_entities = panel[entity_col].nunique()
    T = panel[time_col].nunique()          # assume balanced panel
    N = len(panel)
    k = len(x_cols)

    # ── Step 1: within (FE) residuals to estimate sigma_e² ────────────────────
    result_fe, resid_fe, _, stats_fe, _ = run_within(
        panel, y_col, x_cols, entity_col, time_col
    )
    dof_fe = max(N - k - n_entities, 1)
    sigma_e2 = np.sum(resid_fe ** 2) / dof_fe

    # ── Step 2: between residuals to estimate sigma_u² ────────────────────────
    grp = panel.groupby(entity_col)[[y_col] + list(x_cols)].mean().reset_index()
    y_b  = grp[y_col].values
    X_b  = np.column_stack([np.ones(n_entities)] + [grp[c].values for c in x_cols])
    b_coeffs, _, _, _ = lstsq(X_b, y_b, rcond=None)
    resid_b = y_b - X_b @ b_coeffs
    sigma_b2 = max(0.0, np.sum(resid_b ** 2) / max(n_entities - k - 1, 1) - sigma_e2 / T)
    sigma_u2 = sigma_b2

    # ── Step 3: GLS theta weight ───────────────────────────────────────────────
    theta = 1.0 - np.sqrt(sigma_e2 / max(T * sigma_u2 + sigma_e2, 1e-12))

    # ── Step 4: quasi-demean (partial within) ─────────────────────────────────
    panel2 = panel.copy()
    for col in [y_col] + list(x_cols):
        entity_mean = panel2.groupby(entity_col)[col].transform("mean")
        panel2[col + "_qd"] = panel2[col] - theta * entity_mean

    y_qd = panel2[y_col + "_qd"].values
    X_qd = np.column_stack([np.ones(N)] + [panel2[c + "_qd"].values for c in x_cols])

    # ── Step 5: OLS on quasi-demeaned data ────────────────────────────────────
    coeffs, _, _, _ = lstsq(X_qd, y_qd, rcond=None)
    y_hat  = X_qd @ coeffs
    resid  = y_qd  - y_hat
    dof    = max(N - k - 1, 1)
    s2     = np.sum(resid ** 2) / dof
    cov    = s2 * np.linalg.inv(X_qd.T @ X_qd)
    se     = np.sqrt(np.diag(cov))
    t_stats = coeffs / se
    p_vals  = 2 * sc_stats.t.sf(np.abs(t_stats), df=dof)

    ss_tot = np.sum((y_qd - y_qd.mean()) ** 2)
    ss_res = np.sum(resid ** 2)
    ss_reg = ss_tot - ss_res
    r2     = max(0.0, 1 - ss_res / max(ss_tot, 1e-12))
    r2_adj = max(0.0, 1 - (1 - r2) * (N - 1) / dof)
    f_stat = (ss_reg / max(k, 1)) / (ss_res / dof)
    f_p    = 1 - sc_stats.f.cdf(f_stat, dfn=k, dfd=dof)

    names = ["const"] + list(x_cols)
    result_df = pd.DataFrame({"Variable": names, "Coeff": coeffs, "Std_Err": se,
                               "t_stat": t_stats, "p_value": p_vals})
    stats = {
        "R2": r2, "R2_adj": r2_adj, "N": N, "k": k,
        "AIC": N * np.log(max(ss_res, 1e-10) / N) + 2 * (k + 1),
        "BIC": N * np.log(max(ss_res, 1e-10) / N) + (k + 1) * np.log(N),
        "F_stat": f_stat, "F_p": f_p,
        "sigma_u2": sigma_u2, "sigma_e2": sigma_e2, "theta": theta,
    }
    return result_df, resid, y_hat, stats, cov


def run_fd(df, y_col, x_cols, entity_col, time_col):
    """First-difference estimator."""
    panel = df.sort_values([entity_col, time_col]).copy()
    fd = panel.groupby(entity_col)[[y_col] + list(x_cols)].diff().dropna()
    return run_ols(fd, y_col, x_cols)


def breusch_pagan_test(resid, X):
    """
    Breusch-Pagan / Cook-Weisberg test for heteroskedasticity.
    Regresses squared residuals on the regressors X; the LM statistic
    is n * R² of that auxiliary regression, chi-sq(k) distributed.
    X should include the intercept column.
    """
    from scipy import stats as sc_stats
    from numpy.linalg import lstsq
    resid = np.asarray(resid, dtype=float)
    e2 = resid ** 2
    # auxiliary regression of e² on X
    coeffs_aux, _, _, _ = lstsq(X, e2, rcond=None)
    e2_hat = X @ coeffs_aux
    ss_tot_aux = np.sum((e2 - e2.mean()) ** 2)
    ss_res_aux = np.sum((e2 - e2_hat) ** 2)
    r2_aux = max(0.0, 1 - ss_res_aux / max(ss_tot_aux, 1e-12))
    n = len(resid)
    k = X.shape[1] - 1          # exclude intercept
    bp_stat = n * r2_aux
    bp_p    = 1 - sc_stats.chi2.cdf(bp_stat, df=k)
    return bp_stat, bp_p, k


    """Simple Hausman test statistic."""
    diff = fe_coef - re_coef
    diff_vcov = fe_vcov - re_vcov
    try:
        stat = float(diff @ np.linalg.inv(diff_vcov) @ diff)
        df = len(diff)
        from scipy import stats as sc_stats
        p = 1 - sc_stats.chi2.cdf(stat, df)
        return stat, p, df
    except Exception:
        return None, None, None


def significance_stars(p):
    if p < 0.001: return "***"
    if p < 0.01:  return "**"
    if p < 0.05:  return "*"
    if p < 0.1:   return "·"
    return ""


def call_openai(system_prompt, user_prompt):
    """Call OpenAI GPT-4 API for AI explanations. Key loaded from st.secrets."""
    try:
        api_key = st.secrets["OPENAI_API_KEY"]
    except Exception:
        return "OpenAI API key not configured. Add OPENAI_API_KEY to your Streamlit secrets."

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {api_key}",
    }
    payload = {
        "model": "gpt-4o",
        "max_tokens": 1000,
        "messages": [
            {"role": "system", "content": system_prompt},
            {"role": "user",   "content": user_prompt},
        ],
    }
    try:
        resp = requests.post(
            "https://api.openai.com/v1/chat/completions",
            headers=headers, json=payload, timeout=30,
        )
        data = resp.json()
        if "choices" in data:
            return data["choices"][0]["message"]["content"]
        return f"API error: {data.get('error', {}).get('message', str(data))}"
    except Exception as e:
        return f"Request failed: {e}"


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX REPORT GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════

def build_docx_report(res, model_type, ai_explanation=""):
    """
    Build a professional Word document containing:
      • Cover page
      • Model summary & fit statistics table
      • Coefficient estimates table
      • Residual diagnostics table
      • AI write-up (if available)
    Returns bytes suitable for st.download_button.
    """
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    import datetime, io, numpy as np
    from scipy import stats as sc_stats

    doc = Document()

    # ── Page margins ──────────────────────────────────────────────────────────
    for section in doc.sections:
        section.page_width  = Inches(8.5)
        section.page_height = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)

    # ── Helper: set cell shading ───────────────────────────────────────────────
    def shade_cell(cell, hex_color):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'),   'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'),  hex_color)
        tcPr.append(shd)

    def set_cell_border(cell, **kwargs):
        tc   = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement('w:tcBorders')
        for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            if edge in kwargs:
                tag  = OxmlElement(f'w:{edge}')
                tag.set(qn('w:val'),   kwargs[edge].get('val', 'single'))
                tag.set(qn('w:sz'),    str(kwargs[edge].get('sz', 4)))
                tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
                tcBorders.append(tag)
        tcPr.append(tcBorders)

    # ── Colour palette ────────────────────────────────────────────────────────
    DARK_BG   = "0A0C10"
    ACCENT    = "00C8B0"   # teal (print-safe, slightly darker)
    HEADER_BG = "1A2035"
    ALT_ROW   = "F4F7FA"
    WHITE     = "FFFFFF"
    TEXT_DARK = RGBColor(0x1A, 0x20, 0x35)
    TEAL_RGB  = RGBColor(0x00, 0xC8, 0xB0)
    GRAY_RGB  = RGBColor(0x6B, 0x7A, 0x9A)

    # ── Font helpers ──────────────────────────────────────────────────────────
    def h_style(para, text, size=14, bold=True, color=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=12, space_after=6):
        para.alignment = align
        para.paragraph_format.space_before = Pt(space_before)
        para.paragraph_format.space_after  = Pt(space_after)
        run = para.add_run(text)
        run.bold       = bold
        run.font.size  = Pt(size)
        run.font.name  = "Arial"
        if color:
            run.font.color.rgb = color
        return run

    def body_para(doc, text, size=10, color=None, bold=False, italic=False, space_after=4, align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph()
        p.alignment = align
        p.paragraph_format.space_after = Pt(space_after)
        r = p.add_run(text)
        r.font.size = Pt(size)
        r.font.name = "Arial"
        r.bold      = bold
        r.italic    = italic
        if color:
            r.font.color.rgb = color
        return p

    def add_rule(doc, color_hex="D0D8E8", thickness=6):
        """Thin horizontal rule via paragraph border."""
        p    = doc.add_paragraph()
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    str(thickness))
        bot.set(qn('w:color'), color_hex)
        bot.set(qn('w:space'), '1')
        pBdr.append(bot)
        pPr.append(pBdr)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(6)

    # ── COVER PAGE ────────────────────────────────────────────────────────────
    # Dark-header block via a 1-row table
    cover_tbl = doc.add_table(rows=1, cols=1)
    cover_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    cover_cell = cover_tbl.rows[0].cells[0]
    shade_cell(cover_cell, DARK_BG)
    cover_cell.width = Inches(6.5)
    cover_cell._tc.get_or_add_tcPr()

    cp = cover_cell.add_paragraph()
    cp.paragraph_format.space_before = Pt(18)
    cp.paragraph_format.space_after  = Pt(2)
    cp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r1 = cp.add_run("⬡  PanelStatX")
    r1.font.size = Pt(22)
    r1.font.bold = True
    r1.font.name = "Arial"
    r1.font.color.rgb = TEAL_RGB

    cp2 = cover_cell.add_paragraph()
    cp2.paragraph_format.space_before = Pt(2)
    cp2.paragraph_format.space_after  = Pt(4)
    cp2.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r2 = cp2.add_run("Panel Regression Analysis Report")
    r2.font.size  = Pt(14)
    r2.font.name  = "Arial"
    r2.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF4)

    cp3 = cover_cell.add_paragraph()
    cp3.paragraph_format.space_before = Pt(0)
    cp3.paragraph_format.space_after  = Pt(18)
    cp3.alignment = WD_ALIGN_PARAGRAPH.LEFT
    now_str = datetime.datetime.now().strftime("%d %B %Y, %H:%M")
    r3 = cp3.add_run(f"Generated: {now_str}   ·   Estimator: {model_type}")
    r3.font.size  = Pt(9)
    r3.font.name  = "Arial"
    r3.font.color.rgb = RGBColor(0x6B, 0x7A, 0x9A)

    doc.add_paragraph()  # spacer

    # Cover meta row
    result_df = res["result_df"]
    stats     = res["stats"]
    resid     = np.asarray(res["resid"], dtype=float)
    resid     = resid[np.isfinite(resid)]
    y_col     = res["y_col"]
    x_cols    = res["x_cols"]
    entity_col = res["entity_col"]
    time_col   = res["time_col"]

    meta_items = [
        ("Dependent Variable", y_col),
        ("Independent Variables", ", ".join(x_cols)),
        ("Entity Column", entity_col),
        ("Time Column", time_col),
        ("Observations (N)", f"{stats['N']:,}"),
        ("Variables (k)", str(stats["k"])),
    ]
    meta_tbl = doc.add_table(rows=len(meta_items), cols=2)
    meta_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    col_widths = [Inches(2.2), Inches(4.3)]
    for i, (label, value) in enumerate(meta_items):
        row = meta_tbl.rows[i]
        lc, vc = row.cells[0], row.cells[1]
        lc.width = col_widths[0]
        vc.width = col_widths[1]
        for edge in ('top', 'bottom', 'left', 'right'):
            set_cell_border(lc, **{edge: {'val': 'none', 'sz': 0, 'color': 'FFFFFF'}})
            set_cell_border(vc, **{edge: {'val': 'none', 'sz': 0, 'color': 'FFFFFF'}})
        lp = lc.add_paragraph(label)
        lp.runs[0].font.size = Pt(9)
        lp.runs[0].font.name = "Arial"
        lp.runs[0].bold = True
        lp.runs[0].font.color.rgb = GRAY_RGB
        vp = vc.add_paragraph(value)
        vp.runs[0].font.size = Pt(9)
        vp.runs[0].font.name = "Arial"
        vp.runs[0].font.color.rgb = TEXT_DARK

    doc.add_page_break()

    # ── SECTION 1: MODEL FIT STATISTICS ───────────────────────────────────────
    s1 = doc.add_paragraph()
    h_style(s1, "1.  Model Fit Statistics", size=13, color=TEXT_DARK, space_before=0)
    add_rule(doc)

    fit_rows = [
        ("R²",                f"{stats['R2']:.6f}",          "Proportion of variance explained"),
        ("Adjusted R²",       f"{stats['R2_adj']:.6f}",      "Penalised for number of predictors"),
        ("F-stat",       f"{stats['F_stat']:.6f}",      "Explains regressor outcome"),
         ("p-Value",       f"{stats['F_p']:.6f}",      "Explains model significance"),
        ("Observations (N)",  f"{stats['N']:,}",              "Total data points used in estimation"),
        ("Variables (k)",     f"{stats['k']}",                "Number of independent variables"),
        ("AIC",               f"{stats['AIC']:.4f}",          "Akaike Information Criterion (lower = better)"),
        ("BIC",               f"{stats['BIC']:.4f}",          "Bayesian Information Criterion (lower = better)"),
        ("Estimator",         model_type,                     "Panel regression method applied"),
    ]

    fit_tbl = doc.add_table(rows=len(fit_rows) + 1, cols=3)
    fit_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    fit_tbl.style = "Table Grid"
    fit_col_w = [Inches(2.0), Inches(1.8), Inches(2.7)]

    # Header row
    hdr_cells = fit_tbl.rows[0].cells
    for idx, (hdr_txt, w) in enumerate(zip(["Statistic", "Value", "Description"], fit_col_w)):
        hdr_cells[idx].width = w
        shade_cell(hdr_cells[idx], HEADER_BG)
        p = hdr_cells[idx].paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(hdr_txt)
        r.font.bold  = True
        r.font.size  = Pt(9)
        r.font.name  = "Arial"
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF4)

    for i, (stat_name, stat_val, stat_desc) in enumerate(fit_rows):
        row = fit_tbl.rows[i + 1]
        fill = ALT_ROW if i % 2 == 0 else WHITE
        for j, (txt, w) in enumerate(zip([stat_name, stat_val, stat_desc], fit_col_w)):
            c = row.cells[j]
            c.width = w
            shade_cell(c, fill)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(txt)
            r.font.size = Pt(9)
            r.font.name = "Arial"
            if j == 0:
                r.bold = True
                r.font.color.rgb = TEXT_DARK
            elif j == 1:
                r.font.color.rgb = TEAL_RGB
                r.bold = True
            else:
                r.font.color.rgb = GRAY_RGB

    doc.add_paragraph()

    # ── SECTION 2: COEFFICIENT ESTIMATES TABLE ────────────────────────────────
    s2 = doc.add_paragraph()
    h_style(s2, "2.  Coefficient Estimates", size=13, color=TEXT_DARK, space_before=6)
    add_rule(doc)

    note = body_para(doc, "Significance codes:  *** p<0.001   ** p<0.01   * p<0.05   · p<0.1", size=8, color=GRAY_RGB, italic=True)

    coef_headers = ["Variable", "Coefficient", "Std. Error", "t-Statistic", "p-Value", "Sig.", "95% CI Lower", "95% CI Upper"]
    coef_col_w   = [Inches(1.25), Inches(0.9), Inches(0.8), Inches(0.85), Inches(0.75), Inches(0.4), Inches(0.85), Inches(0.85)]

    coef_tbl = doc.add_table(rows=len(result_df) + 1, cols=len(coef_headers))
    coef_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    coef_tbl.style = "Table Grid"

    # Header
    for j, (h, w) in enumerate(zip(coef_headers, coef_col_w)):
        c = coef_tbl.rows[0].cells[j]
        c.width = w
        shade_cell(c, HEADER_BG)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.font.bold  = True
        r.font.size  = Pt(8)
        r.font.name  = "Arial"
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF4)

    for i, row_data in result_df.iterrows():
        p_val    = row_data["p_value"]
        stars    = significance_stars(p_val)
        ci_lo    = row_data["Coeff"] - 1.96 * row_data["Std_Err"]
        ci_hi    = row_data["Coeff"] + 1.96 * row_data["Std_Err"]
        is_sig   = p_val < 0.05
        fill     = "E8FAF7" if is_sig else (ALT_ROW if i % 2 == 0 else WHITE)

        row_vals = [
            row_data["Variable"],
            f"{row_data['Coeff']:.4f}",
            f"{row_data['Std_Err']:.4f}",
            f"{row_data['t_stat']:.3f}",
            f"{p_val:.4f}",
            stars if stars else "—",
            f"{ci_lo:.4f}",
            f"{ci_hi:.4f}",
        ]
        tbl_row = coef_tbl.rows[i + 1]
        for j, (val, w) in enumerate(zip(row_vals, coef_col_w)):
            c = tbl_row.cells[j]
            c.width = w
            shade_cell(c, fill)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(val)
            r.font.size = Pt(8.5)
            r.font.name = "Arial"
            if j == 0:
                r.bold = True
                r.font.color.rgb = TEXT_DARK
            elif j == 1 and is_sig:
                r.font.color.rgb = TEAL_RGB
                r.bold = True
            elif j == 5 and stars:
                r.font.color.rgb = TEAL_RGB
                r.bold = True
            else:
                r.font.color.rgb = TEXT_DARK

    doc.add_paragraph()

    # ── SECTION 3: RESIDUAL DIAGNOSTICS ───────────────────────────────────────
    doc.add_page_break()
    s3 = doc.add_paragraph()
    h_style(s3, "3.  Residual Diagnostics", size=13, color=TEXT_DARK, space_before=0)
    add_rule(doc)

    jb_stat, jb_p = sc_stats.jarque_bera(resid)
    dw_stat        = np.sum(np.diff(resid)**2) / max(np.sum(resid**2), 1e-10)
    skewness       = sc_stats.skew(resid)
    kurt           = sc_stats.kurtosis(resid)

    # Interpretations
    normality_note  = "Fail to reject normality (p ≥ 0.05)" if jb_p >= 0.05 else "Reject normality — consider robust SEs (p < 0.05)"
    dw_note         = "No autocorrelation detected" if 1.5 <= dw_stat <= 2.5 else f"Possible autocorrelation (DW = {dw_stat:.3f})"
    skew_note       = "Approximately symmetric" if abs(skewness) < 0.5 else ("Moderately skewed" if abs(skewness) < 1 else "Highly skewed")
    kurt_note       = "Approximately normal kurtosis" if abs(kurt) < 1 else ("Leptokurtic (heavy tails)" if kurt > 1 else "Platykurtic (light tails)")

    # Extract Breusch-Pagan results from stats (may be None for some estimators)
    bp_stat = stats.get("BP_stat")
    bp_p    = stats.get("BP_p")
    bp_stat_str = f"{bp_stat:.4f}" if bp_stat is not None else "N/A"
    bp_p_str    = f"{bp_p:.4f}"    if bp_p    is not None else "N/A"
    bp_p_note   = ("p-value >0.05 confirm absence of heteroskedasticity"
                   if bp_p is not None else "Could not be computed for this estimator/data combination")

    diag_rows = [
        ("Mean Residual",       f"{np.mean(resid):.6f}",    "Should be near zero for unbiased model"),
        ("Std. Dev. Residual",  f"{np.std(resid):.6f}",     "Spread of residuals around zero"),
        ("Min Residual",        f"{np.min(resid):.6f}",     "Largest negative deviation"),
        ("Max Residual",        f"{np.max(resid):.6f}",     "Largest positive deviation"),
        ("Skewness",            f"{skewness:.4f}",           skew_note),
        ("Excess Kurtosis",     f"{kurt:.4f}",               kurt_note),
        ("Breusch-Pagan Test",      bp_stat_str,             "Breusch-Pagan Lagrange Multiplier"),
        ("Breusch-Pagan p-Value",   bp_p_str,                bp_p_note),
        ("Jarque-Bera Statistic", f"{jb_stat:.4f}",         "Tests for normality of residuals"),
        ("Jarque-Bera p-value", f"{jb_p:.4f}",              normality_note),
        ("Durbin-Watson",       f"{dw_stat:.4f}",            dw_note),
    ]

    diag_tbl = doc.add_table(rows=len(diag_rows) + 1, cols=3)
    diag_tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
    diag_tbl.style = "Table Grid"
    diag_col_w = [Inches(2.2), Inches(1.5), Inches(2.8)]

    for j, (h, w) in enumerate(zip(["Diagnostic", "Value", "Interpretation"], diag_col_w)):
        c = diag_tbl.rows[0].cells[j]
        c.width = w
        shade_cell(c, HEADER_BG)
        p = c.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        r = p.add_run(h)
        r.font.bold  = True
        r.font.size  = Pt(9)
        r.font.name  = "Arial"
        r.font.color.rgb = RGBColor(0xE2, 0xE8, 0xF4)

    # Flag rows that indicate issues
    concern_rows = {7, 8}  # JB p-value and DW rows — colour-code if concern
    for i, (d_name, d_val, d_note) in enumerate(diag_rows):
        row = diag_tbl.rows[i + 1]
        fill = ALT_ROW if i % 2 == 0 else WHITE
        # Highlight concern rows in amber tint
        if i == 7 and jb_p < 0.05:
            fill = "FFF8E8"
        if i == 8 and (dw_stat < 1.5 or dw_stat > 2.5):
            fill = "FFF8E8"

        for j, (txt, w) in enumerate(zip([d_name, d_val, d_note], diag_col_w)):
            c = row.cells[j]
            c.width = w
            shade_cell(c, fill)
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            r = p.add_run(txt)
            r.font.size = Pt(9)
            r.font.name = "Arial"
            if j == 0:
                r.bold = True
                r.font.color.rgb = TEXT_DARK
            elif j == 1:
                r.font.color.rgb = TEAL_RGB
                r.bold = True
            else:
                r.font.color.rgb = GRAY_RGB

    doc.add_paragraph()

    # ── SECTION 4: AI WRITE-UP ─────────────────────────────────────────────────
    if ai_explanation and ai_explanation.strip():
        doc.add_page_break()
        s4 = doc.add_paragraph()
        h_style(s4, "4.  AI Interpretation & Recommendations", size=13, color=TEXT_DARK, space_before=0)
        add_rule(doc)

        body_para(doc,
            "The following interpretation was generated by GPT-4 based on the regression output above.",
            size=8, color=GRAY_RGB, italic=True, space_after=8)

        # Render the AI text preserving paragraph breaks
        for block in ai_explanation.strip().split("\n"):
            block = block.strip()
            if not block:
                doc.add_paragraph().paragraph_format.space_after = Pt(2)
                continue
            p = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(5)
            p.paragraph_format.left_indent  = Inches(0.15)
            r = p.add_run(block)
            r.font.size = Pt(9.5)
            r.font.name = "Arial"
            r.font.color.rgb = TEXT_DARK
    else:
        s4 = doc.add_paragraph()
        h_style(s4, "4.  AI Interpretation & Recommendations", size=13, color=TEXT_DARK, space_before=6)
        add_rule(doc)
        body_para(doc,
            "No AI interpretation has been generated for this session. "
            "Run the AI Explainer from the AI Explainer tab, then download the report again to include it.",
            size=9, color=GRAY_RGB, italic=True)

    # ── FOOTER on every page ──────────────────────────────────────────────────
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        fp.clear()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        fr = fp.add_run(f"PanelStatX  ·  Panel Regression Analysis Report  ·  {datetime.datetime.now().strftime('%Y-%m-%d')}")
        fr.font.size = Pt(7)
        fr.font.name = "Arial"
        fr.font.color.rgb = GRAY_RGB

    # ── Serialise to bytes ────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ═══════════════════════════════════════════════════════════════════════════════
# SESSION STATE
# ═══════════════════════════════════════════════════════════════════════════════

for key, default in [
    ("df", None), ("results", None), ("ai_explanation", ""),
    ("model_type", "Fixed Effects (Two-Way)"),
    ("access_granted", False), ("access_error", ""),
    # Credit-system state
    ("user_key", ""), ("user_credits", 0),
    ("user_email", ""), ("user_row", None),
]:
    if key not in st.session_state:
        st.session_state[key] = default

# ═══════════════════════════════════════════════════════════════════════════════
# ACCESS KEY GATE  — Google Sheets credit system
# secrets.toml must contain:
#   GSHEET_ID = "your-spreadsheet-id"
#   [gcp_service_account]   ← paste your service-account JSON fields here
#   type = "service_account"
#   project_id = "…"
#   private_key_id = "…"
#   private_key = "-----BEGIN RSA PRIVATE KEY-----\n…\n-----END RSA PRIVATE KEY-----\n"
#   client_email = "…"
#   … (remaining fields from the JSON key file)
# ═══════════════════════════════════════════════════════════════════════════════

if not st.session_state.access_granted:
    st.markdown("""
    <style>
    @import url('https://fonts.googleapis.com/css2?family=DM+Mono:wght@300;400;500&family=Syne:wght@400;600;700;800&display=swap');

    [data-testid="stSidebar"] { display: none !important; }
    [data-testid="stSidebarCollapsedControl"] { display: none !important; }
    [data-testid="stAppViewContainer"] { padding: 0 !important; }
    [data-testid="block-container"] { padding: 0 !important; max-width: 100% !important; }
    section.main > div { padding: 0 !important; }

    /* ── Landing Shell ── */
    .landing-root {
        min-height: 100vh;
        background: #07090e;
        display: flex;
        flex-direction: column;
        position: relative;
        overflow: hidden;
    }

    /* ── Animated grid background ── */
    .grid-bg {
        position: fixed;
        inset: 0;
        background-image:
            linear-gradient(rgba(0,229,200,0.04) 1px, transparent 1px),
            linear-gradient(90deg, rgba(0,229,200,0.04) 1px, transparent 1px);
        background-size: 48px 48px;
        animation: gridPan 24s linear infinite;
        z-index: 0;
    }
    @keyframes gridPan {
        0%   { background-position: 0 0; }
        100% { background-position: 48px 48px; }
    }

    /* ── Glow orbs ── */
    .orb {
        position: fixed;
        border-radius: 50%;
        filter: blur(90px);
        opacity: 0.18;
        z-index: 0;
        animation: orbFloat 10s ease-in-out infinite alternate;
    }
    .orb-1 { width: 520px; height: 520px; background: #00e5c8; top: -160px; left: -160px; animation-delay: 0s; }
    .orb-2 { width: 400px; height: 400px; background: #7c6df0; bottom: -100px; right: -100px; animation-delay: -5s; }
    .orb-3 { width: 280px; height: 280px; background: #f05c7c; top: 40%; left: 55%; animation-delay: -3s; }
    @keyframes orbFloat {
        from { transform: translate(0, 0) scale(1); }
        to   { transform: translate(30px, -20px) scale(1.08); }
    }

    /* ── Top nav bar ── */
    .landing-nav {
        position: relative; z-index: 10;
        display: flex; align-items: center; justify-content: space-between;
        padding: 22px 48px;
        border-bottom: 1px solid rgba(255,255,255,0.05);
        backdrop-filter: blur(12px);
        background: rgba(7,9,14,0.6);
    }
    .nav-logo {
        display: flex; align-items: center; gap: 10px;
        font-family: 'Syne', sans-serif; font-size: 1.25rem; font-weight: 800;
        color: #e2e8f4; letter-spacing: -0.02em;
    }
    .nav-logo .hex { color: #00e5c8; font-size: 1.5rem; }
    .nav-logo .acc { color: #00e5c8; }
    .nav-tag {
        font-family: 'DM Mono', monospace; font-size: 0.62rem;
        color: #6b7a9a; letter-spacing: 0.12em; text-transform: uppercase;
        border: 1px solid rgba(0,229,200,0.2); padding: 4px 12px; border-radius: 20px;
    }
    .nav-right { display: flex; align-items: center; gap: 16px; }
    .nav-badge {
        font-family: 'DM Mono', monospace; font-size: 0.6rem; letter-spacing: 0.1em;
        text-transform: uppercase; padding: 4px 12px; border-radius: 20px;
        background: rgba(0,229,200,0.08); color: #00e5c8;
        border: 1px solid rgba(0,229,200,0.2);
    }

    /* ── Main two-column layout ── */
    .landing-body {
        position: relative; z-index: 5;
        display: flex; align-items: center; justify-content: center;
        gap: 80px; padding: 64px 48px;
        flex: 1;
    }

    /* ── Left hero column ── */
    .hero-col { flex: 1.1; max-width: 560px; }

    .hero-eyebrow {
        display: inline-flex; align-items: center; gap: 8px;
        font-family: 'DM Mono', monospace; font-size: 0.65rem;
        letter-spacing: 0.16em; text-transform: uppercase; color: #00e5c8;
        margin-bottom: 24px;
    }
    .eyebrow-dot {
        width: 6px; height: 6px; border-radius: 50%;
        background: #00e5c8;
        box-shadow: 0 0 8px #00e5c8;
        animation: pulse 2s ease-in-out infinite;
    }
    @keyframes pulse { 0%,100%{opacity:1;transform:scale(1)} 50%{opacity:0.4;transform:scale(0.7)} }

    .hero-title {
        font-family: 'Syne', sans-serif; font-weight: 800;
        font-size: clamp(2.6rem, 4vw, 3.6rem);
        line-height: 1.0; letter-spacing: -0.03em;
        color: #e2e8f4; margin: 0 0 20px 0;
    }
    .hero-title .acc  { color: #00e5c8; }
    .hero-title .acc2 { color: #7c6df0; }

    .hero-desc {
        font-family: 'DM Mono', monospace; font-size: 0.82rem;
        color: #8a96b0; line-height: 1.8; margin-bottom: 36px;
        max-width: 460px;
    }

    /* ── Feature list ── */
    .feature-list { display: flex; flex-direction: column; gap: 12px; margin-bottom: 40px; }
    .feature-item {
        display: flex; align-items: flex-start; gap: 14px;
        font-family: 'DM Mono', monospace; font-size: 0.78rem; color: #8a96b0;
    }
    .feature-icon {
        width: 28px; height: 28px; border-radius: 6px; flex-shrink: 0;
        display: flex; align-items: center; justify-content: center;
        font-size: 0.85rem;
    }
    .fi-teal   { background: rgba(0,229,200,0.1); border: 1px solid rgba(0,229,200,0.25); }
    .fi-purple { background: rgba(124,109,240,0.1); border: 1px solid rgba(124,109,240,0.25); }
    .fi-pink   { background: rgba(240,92,124,0.1); border: 1px solid rgba(240,92,124,0.25); }
    .fi-amber  { background: rgba(245,166,35,0.1); border: 1px solid rgba(245,166,35,0.25); }
    .feature-text b { color: #c8d0e0; }

    /* ── Stat row ── */
    .stat-row { display: flex; gap: 28px; }
    .stat-chip {
        background: rgba(255,255,255,0.03);
        border: 1px solid rgba(255,255,255,0.07);
        border-radius: 8px; padding: 14px 20px;
        font-family: 'DM Mono', monospace;
    }
    .stat-chip .sv {
        font-family: 'Syne', sans-serif; font-size: 1.5rem; font-weight: 700;
        color: #00e5c8; line-height: 1;
    }
    .stat-chip .sl { font-size: 0.62rem; color: #6b7a9a; text-transform: uppercase; letter-spacing: 0.1em; margin-top: 4px; }

    /* ── Access card ── */
    .gate-col { width: 380px; flex-shrink: 0; }

    .gate-card {
        background: rgba(17,19,24,0.85);
        backdrop-filter: blur(24px);
        border: 1px solid rgba(255,255,255,0.08);
        border-radius: 16px;
        padding: 36px 32px 28px;
        box-shadow:
            0 0 0 1px rgba(0,229,200,0.06),
            0 32px 64px rgba(0,0,0,0.6),
            inset 0 1px 0 rgba(255,255,255,0.06);
        position: relative;
        overflow: hidden;
    }
    .gate-card::before {
        content: '';
        position: absolute; top: 0; left: 0; right: 0; height: 2px;
        background: linear-gradient(90deg, transparent, #00e5c8, #7c6df0, transparent);
    }

    .gate-header { margin-bottom: 28px; }
    .gate-title {
        font-family: 'Syne', sans-serif; font-size: 1.35rem; font-weight: 700;
        color: #e2e8f4; letter-spacing: -0.02em; margin: 0 0 6px 0;
    }
    .gate-sub {
        font-family: 'DM Mono', monospace; font-size: 0.68rem;
        color: #6b7a9a; letter-spacing: 0.06em;
    }

    .gate-label {
        font-family: 'DM Mono', monospace; font-size: 0.62rem;
        text-transform: uppercase; letter-spacing: 0.12em;
        color: #6b7a9a; margin-bottom: 8px;
    }

    .gate-footer {
        margin-top: 20px; text-align: center;
        font-family: 'DM Mono', monospace; font-size: 0.6rem;
        color: #4a5568; line-height: 1.7;
    }
    .gate-footer a { color: #00e5c8; text-decoration: none; }

    .error-box {
        margin-top: 12px; padding: 10px 14px;
        background: rgba(240,92,124,0.07);
        border: 1px solid rgba(240,92,124,0.25); border-radius: 6px;
        font-family: 'DM Mono', monospace; font-size: 0.7rem; color: #f05c7c;
        display: flex; align-items: center; gap: 8px;
    }

    /* ── Ticker strip ── */
    .ticker-strip {
        position: relative; z-index: 5;
        border-top: 1px solid rgba(255,255,255,0.05);
        background: rgba(7,9,14,0.7);
        backdrop-filter: blur(12px);
        overflow: hidden; padding: 12px 0;
        white-space: nowrap;
    }
    .ticker-inner {
        display: inline-flex; gap: 0;
        animation: ticker 28s linear infinite;
    }
    @keyframes ticker { from { transform: translateX(0); } to { transform: translateX(-50%); } }
    .ticker-item {
        font-family: 'DM Mono', monospace; font-size: 0.65rem;
        color: #4a5568; letter-spacing: 0.08em; padding: 0 40px;
    }
    .ticker-item span { color: #00e5c8; margin-right: 8px; }

    /* ── Stagger-in animations ── */
    .fade-up { opacity: 0; transform: translateY(18px); animation: fadeUp 0.6s ease forwards; }
    @keyframes fadeUp { to { opacity: 1; transform: translateY(0); } }
    .d1{animation-delay:0.05s} .d2{animation-delay:0.15s} .d3{animation-delay:0.25s}
    .d4{animation-delay:0.35s} .d5{animation-delay:0.45s} .d6{animation-delay:0.55s}
    </style>
    """, unsafe_allow_html=True)

    # ── Background layers ──────────────────────────────────────────────────────
    st.markdown("""
    <div class="grid-bg"></div>
    <div class="orb orb-1"></div>
    <div class="orb orb-2"></div>
    <div class="orb orb-3"></div>
    """, unsafe_allow_html=True)

    # ── Nav bar ───────────────────────────────────────────────────────────────
    st.markdown("""
    <div class="landing-nav fade-up d1">
        <div class="nav-logo">
            <span class="hex">⬡</span>
            PanelStatX<span class="acc"> 
        </div>
        <span class="nav-tag">Panel Regression Engine</span>
    </div>
    """, unsafe_allow_html=True)

    # ── Two-column body ───────────────────────────────────────────────────────
    left_col, right_col = st.columns([1.25, 1], gap="large")

    with left_col:
        st.markdown("""
        <div style="padding: 32px 0 24px 0;">

          <div class="hero-eyebrow fade-up d2">
            <span class="eyebrow-dot"></span>
            Advanced Econometrics Platform
          </div>

          <h1 class="hero-title fade-up d2">
            Panel Data<br>
            <span class="acc">Analysis</span><br>
            <span class="acc2">Re-imagined</span>
          </h1>

          <p class="hero-desc fade-up d3">
            Production-grade panel regression with Fixed Effects, Random Effects,
            First-Difference estimators and AI-powered interpretation — built for
            economists, researchers, and data scientists who demand rigour.
          </p>

          <div class="feature-list fade-up d4">
            <div class="feature-item">
              <div class="feature-icon fi-teal">⬡</div>
              <div class="feature-text"><b>Multiple Estimators</b> — OLS, Fixed Effects (Two-Way), Random Effects, First-Difference with Hausman test</div>
            </div>
            <div class="feature-item">
              <div class="feature-icon fi-purple">◈</div>
              <div class="feature-text"><b>AI Explainer</b> — GPT-4 powered coefficient interpretation with economic significance analysis</div>
            </div>
            <div class="feature-item">
              <div class="feature-icon fi-pink">◉</div>
              <div class="feature-text"><b>Diagnostics Suite</b> — Residual plots, Q-Q, Jarque-Bera normality, Durbin-Watson autocorrelation</div>
            </div>
            <div class="feature-item">
              <div class="feature-icon fi-amber">◆</div>
              <div class="feature-text"><b>Entity Visualisation</b> — Cross-sectional panel plots, time-series overlays, mean comparisons</div>
            </div>
          </div>

          <div class="stat-row fade-up d5">
            <div class="stat-chip">
              <div class="sv">4+</div>
              <div class="sl">Estimators</div>
            </div>
            <div class="stat-chip">
              <div class="sv">CSV · XLS</div>
              <div class="sl">Data Formats</div>
            </div>
            <div class="stat-chip">
              <div class="sv">GPT-4</div>
              <div class="sl">AI Engine</div>
            </div>
          </div>
        </div>
        """, unsafe_allow_html=True)

    with right_col:
        st.markdown("""
        <div style="display:flex; align-items:center; justify-content:center; padding: 32px 0 24px 0;">
        <div class="gate-card fade-up d3" style="width:100%;">
            <div class="gate-header">
                <div class="gate-title">🔐Access Key Required</div>
                <div class="gate-sub">Enter your licence key to continue </div>
            </div>
            <div class="gate-label">Access Key format: PSX-****-****-****</div>
        </div>
        </div>
        """, unsafe_allow_html=True)

        # ── Streamlit input widgets (native, inside styled card) ────────────
        st.markdown("""
        <style>
        /* Tighten up the input so it sits flush inside the card look */
        [data-testid="stTextInput"] > div > div > input {
            background: rgba(255,255,255,0.04) !important;
            border: 1px solid rgba(0,229,200,0.2) !important;
            border-radius: 8px !important;
            color: #e2e8f4 !important;
            font-family: 'DM Mono', monospace !important;
            font-size: 0.82rem !important;
            padding: 12px 16px !important;
            transition: border-color 0.2s !important;
        }
        [data-testid="stTextInput"] > div > div > input:focus {
            border-color: rgba(0,229,200,0.6) !important;
            box-shadow: 0 0 0 3px rgba(0,229,200,0.08) !important;
        }
        [data-testid="stTextInput"] > div > div > input::placeholder { color: #4a5568 !important; }

        /* Primary button - override to landing style */
        [data-testid="baseButton-primary"] {
            background: linear-gradient(135deg, #00e5c8, #00c4ab) !important;
            border: none !important;
            color: #07090e !important;
            font-family: 'Syne', sans-serif !important;
            font-weight: 700 !important;
            font-size: 0.85rem !important;
            letter-spacing: 0.04em !important;
            border-radius: 8px !important;
            padding: 14px !important;
            transition: all 0.2s !important;
            box-shadow: 0 4px 20px rgba(0,229,200,0.25) !important;
        }
        [data-testid="baseButton-primary"]:hover {
            box-shadow: 0 6px 28px rgba(0,229,200,0.4) !important;
            transform: translateY(-1px) !important;
        }
        </style>
        """, unsafe_allow_html=True)

        entered_key = st.text_input(
            "Access Key",
            type="password",
            placeholder="PSX-XXXX-XXXX-XXXX",
            label_visibility="collapsed",
        )
        unlock_btn = st.button("Unlock PanelStatX 🔑", use_container_width=True, type="primary")

        if st.session_state.access_error:
            st.markdown(f"""
            <div class="error-box">
                <span>✕</span> {st.session_state.access_error}
            </div>
            """, unsafe_allow_html=True)


        st.markdown("""
        <div class="gate-footer">
            Access is licence-controlled.<br>
            Contact Administrator  <a href="https://wa.me/2348096506034"> Here →</a> <br>
            
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown("""
        <div style="padding: 20px 0 16px 0;">

          <div class="hero-eyebrow fade-up d2">
            <span class="eyebrow-dot"></span>
            PAYMENT PLANS
          </div>
          <div class="stat-row fade-up d5">
          
        <a href="FLUTTERWAVE_LINK_12" target="_blank" style="text-decoration:none; flex:1;">
            <div class="stat-chip">
              <div class="sv">$10/5 Credits</div>
              <div class="sl">5 Analysis Run</div>
            </div>
        </a>
        
        <a href="FLUTTERWAVE_LINK_12" target="_blank" style="text-decoration:none; flex:1;">
            <div class="stat-chip">
              <div class="sv">$15/25 Credits</div>
              <div class="sl">25 Analysis Runs</div>
            </div>
        </a>
        
        <a href="FLUTTERWAVE_LINK_12" target="_blank" style="text-decoration:none; flex:1;">
            <div class="stat-chip">
              <div class="sv">$45/30 Credits </div>
              <div class="sl">30 Analysis Runs</div>
            </div>
        </a>
        
          </div>
        </div>
        
        <div class="gate-footer">
           Request for <a href="https://wa.me/2348096506034">🔑Access key →</a> Here
        </div>
        
        """, unsafe_allow_html=True)
        


    # ── Ticker strip ──────────────────────────────────────────────────────────
    ticker_items = [
        ("◈", "Fixed Effects Estimation"),
        ("◉", "Random Effects (GLS)"),
        ("◆", "First-Difference Estimator"),
        ("⬡", "Hausman Specification Test"),
        ("✦", "AI Coefficient Interpretation"),
        ("◈", "Residual Diagnostics"),
        ("◉", "Q-Q Normality Plots"),
        ("◆", "Durbin-Watson Autocorrelation"),
        ("⬡", "Entity Cross-Section Plots"),
        ("✦", "Balanced Panel Support"),
    ]
    ticker_html = "".join(f'<span class="ticker-item"><span>{icon}</span>{label}</span>' for icon, label in ticker_items)
    st.markdown(f"""
    <div class="ticker-strip">
        <div class="ticker-inner">
            {ticker_html}{ticker_html}
        </div>
    </div>
    """, unsafe_allow_html=True)

    if unlock_btn:
        if not entered_key:
            st.session_state.access_error = "Please enter your access key."
            st.rerun()
        else:
            with st.spinner("Verifying key…"):
                record = lookup_key(entered_key)
            if record is None:
                st.session_state.access_error = "Invalid access key. Please try again."
                st.rerun()
            elif record["credits"] <= 0:
                st.session_state.access_error = (
                    "Your credits have been exhausted. "
                    "Please purchase more credits to continue."
                )
                st.rerun()
            else:
                # Valid key with credits — grant access and cache user state
                st.session_state.access_granted = True
                st.session_state.access_error = ""
                st.session_state.user_key = record["key"]
                st.session_state.user_credits = record["credits"]
                st.session_state.user_email = record["email"]
                st.session_state.user_row = record["row_index"]
                st.rerun()

    st.stop()


# ═══════════════════════════════════════════════════════════════════════════════
# SIDEBAR
# ═══════════════════════════════════════════════════════════════════════════════

with st.sidebar:
    st.markdown("""
    <div style="padding:16px 0 24px 0; border-bottom:1px solid var(--border); margin-bottom:20px;">
        <div style="font-family:'Syne',sans-serif;font-size:1.4rem;font-weight:800;color:var(--text);letter-spacing:-0.02em;">
            &#x2B21; Panel<span style="color:var(--accent);">Stat</span>X
        </div>
        <div style="font-family:'DM Mono',monospace;font-size:0.68rem;color:var(--muted);margin-top:4px;letter-spacing:0.08em;">
            PANEL REGRESSION ENGINE v1.0
        </div>
    </div>
    """, unsafe_allow_html=True)

    # ── Credit HUD ────────────────────────────────────────────────────────────
    credits_left = st.session_state.user_credits
    credit_color = "#00e5c8" if credits_left > 5 else "#f5a623" if credits_left > 1 else "#f05c7c"
    credit_label = "Credits remaining" if credits_left > 1 else ("1 credit left!" if credits_left == 1 else "No credits")
    email_display = st.session_state.user_email or "—"
    st.markdown(f"""
    <div style="background:var(--surface2);border:1px solid var(--border);border-radius:8px;
                padding:14px 16px;margin-bottom:20px;font-family:'DM Mono',monospace;">
        <div style="font-size:0.6rem;text-transform:uppercase;letter-spacing:0.12em;
                    color:var(--muted);margin-bottom:8px;">Account</div>
        <div style="font-size:0.7rem;color:var(--muted);margin-bottom:10px;
                    overflow:hidden;text-overflow:ellipsis;white-space:nowrap;">{email_display}</div>
        <div style="display:flex;align-items:baseline;gap:6px;">
            <span style="font-family:'Syne',sans-serif;font-size:1.6rem;font-weight:700;
                         color:{credit_color};line-height:1;">{credits_left}</span>
            <span style="font-size:0.62rem;color:var(--muted);text-transform:uppercase;
                         letter-spacing:0.1em;">{credit_label}</span>
        </div>
        <div style="margin-top:8px;background:var(--border);border-radius:2px;height:3px;">
            <div style="height:3px;border-radius:2px;background:{credit_color};
                        width:{min(100, credits_left * 10)}%;transition:width 0.4s;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

    if credits_left <= 0:
        st.error("⚠ No credits remaining. Please purchase more to run analyses.")
        st.markdown("""
        <div style="text-align:center;font-family:'DM Mono',monospace;font-size:0.7rem;
                    color:var(--muted);padding:8px 0;">
            <a href="#" style="color:var(--accent);text-decoration:none;">Buy more credits →</a>
        </div>
        """, unsafe_allow_html=True)

    st.markdown('<div style="font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Data Source</div>', unsafe_allow_html=True)
    data_src = st.radio("", ["Use Demo Dataset", "Upload File"], label_visibility="collapsed")

    if data_src == "Upload File":
        uploaded = st.file_uploader(
            "Upload panel data",
            type=["csv", "xlsx", "xls"],
            label_visibility="collapsed",
            help="Accepts CSV (.csv) or Excel (.xlsx / .xls) files",
        )
        if uploaded:
            try:
                fname = uploaded.name.lower()
                if fname.endswith(".csv"):
                    st.session_state.df = pd.read_csv(uploaded)
                else:
                    # Excel: let user pick sheet if multiple exist
                    xf = pd.ExcelFile(uploaded)
                    if len(xf.sheet_names) > 1:
                        sheet = st.selectbox("Sheet", xf.sheet_names)
                    else:
                        sheet = xf.sheet_names[0]
                    st.session_state.df = pd.read_excel(uploaded, sheet_name=sheet)
                st.success(f"Loaded {st.session_state.df.shape[0]:,} rows × {st.session_state.df.shape[1]} cols")
            except Exception as e:
                st.error(f"Could not read file: {e}")
    else:
        if st.button("Load Demo Data", use_container_width=True):
            st.session_state.df = generate_demo_panel()
            st.success("Demo panel loaded!")

    st.markdown("---")

    if st.session_state.df is not None:
        df = st.session_state.df
        cols = df.columns.tolist()

        st.markdown('<div style="font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Variable Mapping</div>', unsafe_allow_html=True)
        entity_col = st.selectbox("Entity / Panel ID", cols, index=cols.index("entity") if "entity" in cols else 0)
        time_col   = st.selectbox("Time Variable", cols, index=cols.index("year") if "year" in cols else 1)
        y_col      = st.selectbox("Dependent Variable (Y)", [c for c in cols if c not in [entity_col, time_col]],
                                   index=0)
        x_candidates = [c for c in cols if c not in [entity_col, time_col, y_col]]
        x_cols = st.multiselect("Independent Variables (X)", x_candidates, default=x_candidates[:3] if len(x_candidates) >= 3 else x_candidates)

        st.markdown("---")
        st.markdown('<div style="font-size:0.7rem;text-transform:uppercase;letter-spacing:0.12em;color:var(--muted);margin-bottom:8px;">Estimator</div>', unsafe_allow_html=True)
        model_type = st.selectbox("", ["Fixed Effects (Two-Way)", "Fixed Effects (Entity)", "Random Effects (GLS)", "First Difference", "Pooled OLS"], label_visibility="collapsed")
        st.session_state.model_type = model_type

        st.markdown("---")
        run_btn = st.button(
            "Run Analysis",
            use_container_width=True,
            type="primary",
            disabled=(st.session_state.user_credits <= 0),
        )
    else:
        run_btn = False
        entity_col = time_col = y_col = "—"
        x_cols = []
        model_type = "Fixed Effects (Two-Way)"

    # ── New Analysis reset button ───────────────────────────────────────────────
    st.markdown("---")
    if st.button("New Analysis", use_container_width=True):
        st.session_state.df = None
        st.session_state.results = None
        st.session_state.ai_explanation = ""
        st.rerun()


# ═══════════════════════════════════════════════════════════════════════════════
# RUN MODEL
# ═══════════════════════════════════════════════════════════════════════════════

if run_btn and st.session_state.df is not None and x_cols:
    # ── Credit guard ──────────────────────────────────────────────────────────
    if st.session_state.user_credits <= 0:
        st.error("⚠ You have no credits remaining. Please purchase more to run analyses.")
        st.stop()

    df = st.session_state.df
    with st.spinner("Running regression…"):
        try:
            if model_type == "Pooled OLS":
                result_df, resid, y_hat, stats, vcov = run_ols(df, y_col, x_cols)
            elif model_type == "First Difference":
                result_df, resid, y_hat, stats, vcov = run_fd(df, y_col, x_cols, entity_col, time_col)
            elif model_type == "Fixed Effects (Entity)":
                result_df, resid, y_hat, stats, vcov = run_within(df, y_col, x_cols, entity_col, time_col)
            elif model_type == "Random Effects (GLS)":
                result_df, resid, y_hat, stats, vcov = run_re(df, y_col, x_cols, entity_col, time_col)
            else:  # Fixed Effects (Two-Way)
                result_df, resid, y_hat, stats, vcov = run_within(df, y_col, x_cols, entity_col, time_col)

            # ── Breusch-Pagan heteroskedasticity test ─────────────────────────
            # Build regressor matrix for BP (with intercept)
            try:
                _bp_X = np.column_stack([np.ones(len(resid))] +
                                         [df[c].values[:len(resid)] for c in x_cols])
                bp_stat, bp_p, bp_k = breusch_pagan_test(resid, _bp_X)
                stats["BP_stat"] = bp_stat
                stats["BP_p"]    = bp_p
            except Exception:
                stats["BP_stat"] = None
                stats["BP_p"]    = None

            # ── Hausman test (RE vs FE) ────────────────────────────────────────
            hausman_result = None
            if model_type == "Random Effects (GLS)":
                try:
                    _, _, _, _, fe_vcov = run_within(df, y_col, x_cols, entity_col, time_col)
                    fe_res, _, _, _, _  = run_within(df, y_col, x_cols, entity_col, time_col)
                    fe_coef = fe_res["Coeff"].values
                    re_coef = result_df[result_df["Variable"] != "const"]["Coeff"].values
                    # align – both should be len(x_cols)
                    h_stat, h_p, h_df = hausman_test(fe_coef, re_coef, fe_vcov, vcov[1:, 1:])
                    hausman_result = {"stat": h_stat, "p": h_p, "df": h_df}
                except Exception:
                    pass

            st.session_state.results = {
                "result_df": result_df, "resid": resid, "y_hat": y_hat,
                "stats": stats, "y_col": y_col, "x_cols": x_cols,
                "entity_col": entity_col, "time_col": time_col,
                "hausman": hausman_result,
            }
            st.session_state.ai_explanation = ""

            # ── Deduct 1 credit ───────────────────────────────────────────────
            new_credits = deduct_credit(
                st.session_state.user_row,
                st.session_state.user_credits,
            )
            st.session_state.user_credits = new_credits
            if new_credits == 0:
                st.warning("⚠ You just used your last credit. Purchase more to run further analyses.")
            elif new_credits <= 3:
                st.warning(f"⚠ Only {new_credits} credit(s) remaining.")

        except Exception as e:
            st.error(f"Regression error: {e}")


# ═══════════════════════════════════════════════════════════════════════════════
# MAIN LAYOUT
# ═══════════════════════════════════════════════════════════════════════════════

# Hero
st.markdown("""
<div class="hero">
    <div class="hero-title">Panel<span>Stat</span>X</div>
    <div class="hero-sub">⬡ Panel Regression Analysis System · AI-Powered Econometrics</div>
</div>
""", unsafe_allow_html=True)

if st.session_state.df is None:
    # ── Landing state ──────────────────────────────────────────────────────────
    col1, col2, col3 = st.columns(3)
    for col, icon, title, desc in [
        (col1, "⬡", "Panel-Ready", "Fixed effects, first-difference, and pooled OLS estimators built for longitudinal data."),
        (col2, "◈", "Diagnostic Suite", "Residual analysis, heteroskedasticity checks, Hausman test, and entity plots."),
        (col3, "⬟", "AI Explainer", "GPT-4o interprets your regression output in plain language — coefficients, fit, and caveats."),
    ]:
        with col:
            st.markdown(f"""
            <div class="scard" style="text-align:center;padding:32px 20px;">
                <div style="font-size:2rem;margin-bottom:12px;color:var(--accent);">{icon}</div>
                <div style="font-family:'Syne',sans-serif;font-size:1rem;font-weight:700;color:var(--text);margin-bottom:8px;">{title}</div>
                <div style="font-size:0.78rem;color:var(--muted);line-height:1.6;">{desc}</div>
            </div>
            """, unsafe_allow_html=True)

    st.markdown("""
    <div style="text-align:center;margin-top:40px;padding:32px;background:var(--surface2);border:1px dashed var(--border);border-radius:8px;">
        <div style="font-family:'Syne',sans-serif;font-size:1rem;color:var(--muted);">
            ← Load demo data or upload a CSV from the sidebar to begin
        </div>
    </div>
    """, unsafe_allow_html=True)
    st.stop()


# ── Data has been loaded ────────────────────────────────────────────────────
df = st.session_state.df
res = st.session_state.results

# ── Quick dataset stats bar ───────────────────────────────────────────────────
n_e = df[entity_col].nunique() if entity_col in df.columns else "—"
n_t = df[time_col].nunique() if time_col in df.columns else "—"
st.markdown(f"""
<div style="margin-bottom:24px;">
    <span class="stat-pill">Entities <b>{n_e}</b></span>
    <span class="stat-pill">Periods <b>{n_t}</b></span>
    <span class="stat-pill">Observations <b>{len(df):,}</b></span>
    <span class="stat-pill">Estimator <b>{st.session_state.model_type}</b></span>
    <span class="badge badge-teal" style="margin-left:4px;">READY</span>
</div>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════════════════════════
# TABS
# ═══════════════════════════════════════════════════════════════════════════════

tab_data, tab_results, tab_diagnostics, tab_entity, tab_ai = st.tabs([
    "⬡ Data Explorer", "◈ Results", "⬟ Diagnostics", "⬢ Entity Plots", "✦ AI Explainer"
])


# ──────────────────────────────────────────────────────────────────────────────
# TAB 1 · DATA EXPLORER
# ──────────────────────────────────────────────────────────────────────────────

with tab_data:
    c1, c2 = st.columns([2, 1])
    with c1:
        st.markdown('<div class="scard-title">Dataset Preview</div>', unsafe_allow_html=True)
        st.dataframe(df.head(100), use_container_width=True, height=320)
    with c2:
        st.markdown('<div class="scard-title">Summary Statistics</div>', unsafe_allow_html=True)
        st.dataframe(df.describe().round(3), use_container_width=True, height=320)

    st.markdown("---")
    st.markdown('<div class="scard-title">Correlation Heatmap</div>', unsafe_allow_html=True)
    num_cols = df.select_dtypes(include=[np.number]).columns.tolist()
    if len(num_cols) >= 2:
        corr = df[num_cols].corr().round(3)
        fig_corr = go.Figure(go.Heatmap(
            z=corr.values, x=corr.columns, y=corr.index,
            colorscale=[[0, "#f05c7c"], [0.5, "#111318"], [1, "#00e5c8"]],
            zmid=0, text=corr.values.round(2),
            texttemplate="%{text}", showscale=True,
        ))
        fig_corr.update_layout(title="Pearson Correlation Matrix", height=380, **PLOTLY_THEME)
        st.plotly_chart(fig_corr, use_container_width=True)

    # Distribution of Y
    if y_col in df.columns:
        st.markdown('<div class="scard-title" style="margin-top:16px;">Dependent Variable Distribution</div>', unsafe_allow_html=True)
        fig_dist = px.histogram(df, x=y_col, nbins=40, color_discrete_sequence=["#00e5c8"])
        fig_dist.update_layout(title=f"Distribution of {y_col}", height=300, bargap=0.05, **PLOTLY_THEME)
        st.plotly_chart(fig_dist, use_container_width=True)


# ──────────────────────────────────────────────────────────────────────────────
# TAB 2 · RESULTS
# ──────────────────────────────────────────────────────────────────────────────

with tab_results:
    if res is None:
        st.info("Run the analysis from the sidebar to view regression results.")
    else:
        result_df = res["result_df"]
        stats = res["stats"]

        # ── Model fit summary ──────────────────────────────────────────────────
        st.markdown('<div class="scard-title">Model Fit</div>', unsafe_allow_html=True)
        mc = st.columns(8)
        f_label = f"{stats.get('F_stat', 0):.3f}" if stats.get('F_stat') is not None else "—"
        fp_label = f"{stats.get('F_p', 1):.4f}" if stats.get('F_p') is not None else "—"
        for col, label, val in [
            (mc[0], "R²",        f"{stats['R2']:.4f}"),
            (mc[1], "Adj. R²",   f"{stats['R2_adj']:.4f}"),
            (mc[2], "N",         f"{stats['N']:,}"),
            (mc[3], "Variables", f"{stats['k']}"),
            (mc[4], "AIC",       f"{stats['AIC']:.2f}"),
            (mc[5], "BIC",       f"{stats['BIC']:.2f}"),
            (mc[6], "F-stat",    f_label),
            (mc[7], "F p-value", fp_label),
        ]:
            with col:
                st.metric(label, val)

        # F-stat interpretation
        if stats.get("F_p") is not None:
            if stats["F_p"] < 0.05:
                st.success(f"✓ F-statistic ({stats['F_stat']:.3f}) is significant (p={stats['F_p']:.4f}) — regressors jointly explain the outcome.")
            else:
                st.warning(f"⚠ F-statistic ({stats['F_stat']:.3f}) is not significant (p={stats['F_p']:.4f}) — regressors may not jointly explain the outcome.")

        # RE variance components
        if st.session_state.model_type == "Random Effects (GLS)":
            st.markdown("---")
            st.markdown('<div class="scard-title">Random Effects Variance Components</div>', unsafe_allow_html=True)
            rc1, rc2, rc3 = st.columns(3)
            with rc1: st.metric("σ²ᵤ (between)", f"{stats.get('sigma_u2', 0):.6f}")
            with rc2: st.metric("σ²ₑ (within)",  f"{stats.get('sigma_e2', 0):.6f}")
            with rc3: st.metric("θ (GLS weight)", f"{stats.get('theta', 0):.4f}")
            st.caption("θ → 1 means FE dominates; θ → 0 means OLS/pooled dominates.")

        # Hausman test
        hausman_res = res.get("hausman")
        if hausman_res and hausman_res.get("stat") is not None:
            st.markdown("---")
            st.markdown('<div class="scard-title">Hausman Specification Test (RE vs FE)</div>', unsafe_allow_html=True)
            hc1, hc2, hc3 = st.columns(3)
            with hc1: st.metric("χ² statistic", f"{hausman_res['stat']:.4f}")
            with hc2: st.metric("p-value",       f"{hausman_res['p']:.4f}")
            with hc3: st.metric("df",             f"{hausman_res['df']}")
            if hausman_res["p"] < 0.05:
                st.warning("⚠ Hausman test rejects RE (p < 0.05) — Fixed Effects estimator is preferred (endogenous individual effects suspected).")
            else:
                st.success("✓ Hausman test does not reject RE (p ≥ 0.05) — Random Effects estimator is consistent and efficient.")

        st.markdown("---")

        # ── Coefficient table ──────────────────────────────────────────────────
        st.markdown('<div class="scard-title">Coefficient Estimates</div>', unsafe_allow_html=True)
        display = result_df.copy()
        display["Stars"]  = display["p_value"].apply(significance_stars)
        display["Sig"]    = display["p_value"].apply(
            lambda p: "✓ Significant" if p < 0.05 else "✗ Not sig.")
        display = display.rename(columns={
            "Variable": "Variable", "Coeff": "Coeff.",
            "Std_Err": "Std. Err.", "t_stat": "t-stat", "p_value": "p-value"
        })
        st.dataframe(
            display.style
                .format({"Coeff.": "{:.4f}", "Std. Err.": "{:.4f}",
                         "t-stat": "{:.3f}", "p-value": "{:.4f}"})
                .map(lambda v: "color: #00e5c8" if v == "✓ Significant" else "color: #6b7a9a", subset=["Sig"]),
            use_container_width=True, hide_index=True
        )
        st.caption("*p<0.1  **p<0.05  ***p<0.01")

        st.markdown("---")

        # ── Coefficient plot ───────────────────────────────────────────────────
        st.markdown('<div class="scard-title">Coefficient Plot (with 95% CI)</div>', unsafe_allow_html=True)
        rd = res["result_df"]
        ci_lo = rd["Coeff"] - 1.96 * rd["Std_Err"]
        ci_hi = rd["Coeff"] + 1.96 * rd["Std_Err"]
        colors = ["#00e5c8" if p < 0.05 else "#6b7a9a" for p in rd["p_value"]]

        fig_coef = go.Figure()
        fig_coef.add_hline(y=0, line_dash="dash", line_color="#1f2535")
        for pos, (i, row) in enumerate(rd.iterrows()):
            lo, hi = ci_lo.iloc[pos], ci_hi.iloc[pos]
            fig_coef.add_trace(go.Scatter(
                x=[lo, hi], y=[row["Variable"], row["Variable"]],
                mode="lines", line=dict(color="#1f2535", width=2),
                showlegend=False
            ))
        fig_coef.add_trace(go.Scatter(
            x=rd["Coeff"], y=rd["Variable"], mode="markers",
            marker=dict(size=10, color=colors, line=dict(width=1, color="#0a0c10")),
            name="Coefficient", showlegend=False
        ))
        fig_coef.update_layout(title="Coefficients with 95% Confidence Intervals",
                                height=max(280, len(rd) * 55), **PLOTLY_THEME)
        st.plotly_chart(fig_coef, use_container_width=True)

        # ── Download Report ────────────────────────────────────────────────────
        st.markdown("---")
        st.markdown('<div class="scard-title">Export Report</div>', unsafe_allow_html=True)
        st.markdown("""
        <div style="font-family:'DM Mono',monospace;font-size:0.76rem;color:var(--muted);margin-bottom:12px;line-height:1.7;">
            Downloads a fully formatted MS Word (.docx) report containing: model fit statistics,
            coefficient estimates table with significance stars &amp; 95% CI, residual diagnostics table,
            and the AI write-up if already generated in the AI Explainer tab.
        </div>
        """, unsafe_allow_html=True)

        dl_col1, dl_col2 = st.columns([1, 2])
        with dl_col1:
            try:
                import datetime
                docx_bytes = build_docx_report(
                    res,
                    st.session_state.model_type,
                    ai_explanation=st.session_state.get("ai_explanation", ""),
                )
                fname = f"PanelStatX_Report_{datetime.datetime.now().strftime('%Y%m%d_%H%M')}.docx"
                st.download_button(
                    label="⬇  Download Report (.docx)",
                    data=docx_bytes,
                    file_name=fname,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                    type="primary",
                )
            except Exception as e:
                st.error(f"Report generation error: {e}")
        with dl_col2:
            st.markdown("""
            <div style="font-family:'DM Mono',monospace;font-size:0.7rem;color:var(--muted);padding:10px 0;line-height:1.8;">
                ✦ Tip: Generate the AI Explanation first (AI Explainer tab),
                then return here to download the complete report including the write-up.
            </div>
            """, unsafe_allow_html=True)


# ──────────────────────────────────────────────────────────────────────────────
# TAB 3 · DIAGNOSTICS
# ──────────────────────────────────────────────────────────────────────────────

with tab_diagnostics:
    if res is None:
        st.info("Run the analysis first.")
    else:
        resid = res["resid"]
        y_hat = res["y_hat"]

        dc1, dc2 = st.columns(2)

        # Residuals vs Fitted
        with dc1:
            fig_rv = go.Figure()
            fig_rv.add_hline(y=0, line_dash="dash", line_color="#f05c7c", line_width=1)
            fig_rv.add_trace(go.Scatter(
                x=y_hat, y=resid, mode="markers",
                marker=dict(size=5, color="#00e5c8", opacity=0.6),
                name="Residual"
            ))
            fig_rv.update_layout(title="Residuals vs Fitted",
                                   xaxis_title="Fitted", yaxis_title="Residual", height=340, **PLOTLY_THEME)
            st.plotly_chart(fig_rv, use_container_width=True)

        # Q-Q Plot
        with dc2:
            from scipy import stats as sc_stats
            resid_clean = np.asarray(resid, dtype=float)
            resid_clean = resid_clean[np.isfinite(resid_clean)]
            probplot_result = sc_stats.probplot(resid_clean)
            osm, osr = probplot_result[0]
            slope, intercept, _ = probplot_result[1]
            fig_qq = go.Figure()
            fig_qq.add_trace(go.Scatter(x=list(osm), y=list(osr), mode="markers",
                                         marker=dict(size=4, color="#7c6df0", opacity=0.7), name="Residuals"))
            fig_qq.add_trace(go.Scatter(x=[float(min(osm)), float(max(osm))],
                                         y=[slope * float(min(osm)) + intercept, slope * float(max(osm)) + intercept],
                                         mode="lines", line=dict(color="#f05c7c", dash="dash"), name="Normal"))
            fig_qq.update_layout(title="Normal Q-Q Plot",
                                   xaxis_title="Theoretical Quantiles",
                                   yaxis_title="Sample Quantiles", height=340, **PLOTLY_THEME)
            st.plotly_chart(fig_qq, use_container_width=True)

        dc3, dc4 = st.columns(2)

        # Residual distribution
        with dc3:
            fig_rh = px.histogram(x=resid, nbins=40, color_discrete_sequence=["#7c6df0"])
            fig_rh.update_layout(title="Residual Distribution", xaxis_title="Residual",
                                   height=300, bargap=0.05, **PLOTLY_THEME)
            st.plotly_chart(fig_rh, use_container_width=True)

        # Scale-Location
        with dc4:
            fig_sl = go.Figure()
            fig_sl.add_trace(go.Scatter(
                x=y_hat, y=np.sqrt(np.abs(resid)), mode="markers",
                marker=dict(size=5, color="#f5a623", opacity=0.6)
            ))
            fig_sl.update_layout(title="Scale-Location (√|Residual| vs Fitted)",
                                   xaxis_title="Fitted", yaxis_title="√|Residual|", height=300, **PLOTLY_THEME)
            st.plotly_chart(fig_sl, use_container_width=True)

        # Residual stats
        st.markdown("---")
        st.markdown('<div class="scard-title">Residual Diagnostics Summary</div>', unsafe_allow_html=True)
        from scipy import stats as sc_stats
        resid_arr = np.asarray(resid, dtype=float)
        resid_arr = resid_arr[np.isfinite(resid_arr)]
        jb_stat, jb_p = sc_stats.jarque_bera(resid_arr)
        dw = np.sum(np.diff(resid_arr)**2) / max(np.sum(resid_arr**2), 1e-10)

        dc5, dc6, dc7, dc8 = st.columns(4)
        with dc5: st.metric("Mean Residual", f"{np.mean(resid_arr):.4f}")
        with dc6: st.metric("Std Residual", f"{np.std(resid_arr):.4f}")
        with dc7: st.metric("Jarque-Bera p", f"{jb_p:.4f}")
        with dc8: st.metric("Durbin-Watson", f"{dw:.4f}")

        if jb_p < 0.05:
            st.warning("⚠ Jarque-Bera test rejects normality (p < 0.05). Consider robust standard errors.")
        if dw < 1.5 or dw > 2.5:
            st.warning(f"⚠ Durbin-Watson = {dw:.3f} suggests possible autocorrelation.")
        else:
            st.success("✓ Durbin-Watson statistic is in the acceptable range.")

        # ── Breusch-Pagan Heteroskedasticity Test ─────────────────────────────
        st.markdown("---")
        st.markdown('<div class="scard-title">Breusch-Pagan Test for Heteroskedasticity</div>', unsafe_allow_html=True)
        bp_stat = res["stats"].get("BP_stat")
        bp_p    = res["stats"].get("BP_p")
        if bp_stat is not None and bp_p is not None:
            bpc1, bpc2 = st.columns(2)
            with bpc1: st.metric("BP LM Statistic", f"{bp_stat:.4f}")
            with bpc2: st.metric("BP p-value",       f"{bp_p:.4f}")
            if bp_p < 0.05:
                st.warning("⚠ Breusch-Pagan test rejects homoskedasticity (p < 0.05). "
                           "Heteroskedastic errors detected — consider heteroskedasticity-robust (HC) standard errors.")
            else:
                st.success("✓ Breusch-Pagan test does not reject homoskedasticity (p ≥ 0.05).")
            st.caption("H₀: Constant variance (homoskedasticity). LM ~ χ²(k).")
        else:
            st.info("Breusch-Pagan test could not be computed for this estimator/data combination.")


# ──────────────────────────────────────────────────────────────────────────────
# TAB 4 · ENTITY PLOTS
# ──────────────────────────────────────────────────────────────────────────────

with tab_entity:
    if entity_col not in df.columns or time_col not in df.columns:
        st.info("Entity and time columns not set.")
    else:
        y_plot = y_col if y_col in df.columns else df.select_dtypes(np.number).columns[0]

        ec1, ec2 = st.columns([1, 3])
        with ec1:
            entities_avail = sorted(df[entity_col].unique())
            selected_entities = st.multiselect(
                "Select entities to plot",
                entities_avail,
                default=entities_avail[:6] if len(entities_avail) >= 6 else entities_avail
            )
        with ec2:
            x_axis = st.selectbox("X axis", [time_col] + [c for c in df.columns if c not in [entity_col]], index=0)

        if selected_entities:
            plot_df = df[df[entity_col].isin(selected_entities)]

            fig_ep = px.line(
                plot_df, x=x_axis, y=y_plot, color=entity_col,
                markers=True,
                color_discrete_sequence=["#00e5c8", "#7c6df0", "#f05c7c", "#f5a623",
                                          "#22d3a0", "#60a5fa", "#fb923c", "#a78bfa"],
            )
            fig_ep.update_layout(title=f"{y_plot} over {x_axis} by {entity_col}",
                                   height=440, **PLOTLY_THEME)
            st.plotly_chart(fig_ep, use_container_width=True)

            # Entity mean bar
            means = df.groupby(entity_col)[y_plot].mean().sort_values(ascending=False)
            fig_bar = px.bar(
                x=means.index, y=means.values,
                color=means.values,
                color_continuous_scale=["#111318", "#00e5c8"],
                labels={"x": entity_col, "y": f"Mean {y_plot}"},
            )
            fig_bar.update_layout(title=f"Entity Mean of {y_plot}", height=320,
                                   coloraxis_showscale=False, **PLOTLY_THEME)
            st.plotly_chart(fig_bar, use_container_width=True)


# ──────────────────────────────────────────────────────────────────────────────
# TAB 5 · AI EXPLAINER
# ──────────────────────────────────────────────────────────────────────────────

with tab_ai:
    st.markdown('<div class="scard-title">AI Regression Explainer</div>', unsafe_allow_html=True)

    if res is None:
        st.info("Run the analysis first to unlock the AI explainer.")
    else:
        result_df = res["result_df"]
        stats = res["stats"]

        # Build a rich summary for GPT-4o Model
        coeff_table = result_df.to_string(index=False)
        hausman_res = res.get("hausman")
        hausman_str = ""
        if hausman_res and hausman_res.get("stat") is not None:
            hausman_str = f"\nHausman Test: χ²={hausman_res['stat']:.4f}, p={hausman_res['p']:.4f} (df={hausman_res['df']})"
        bp_str = ""
        if stats.get("BP_stat") is not None:
            bp_str = f"\nBreusch-Pagan Test: LM={stats['BP_stat']:.4f}, p={stats['BP_p']:.4f}"
        re_str = ""
        if st.session_state.model_type == "Random Effects (GLS)":
            re_str = f"\nVariance Components: σ²ᵤ={stats.get('sigma_u2',0):.6f}, σ²ₑ={stats.get('sigma_e2',0):.6f}, θ={stats.get('theta',0):.4f}"
        context = f"""
Model: {st.session_state.model_type}
Dependent variable: {res['y_col']}
Independent variables: {', '.join(res['x_cols'])}
Entity column: {res['entity_col']} | Time column: {res['time_col']}

Fit Statistics:
  R²        = {stats['R2']:.4f}
  Adj. R²   = {stats['R2_adj']:.4f}
  N         = {stats['N']}
  AIC       = {stats['AIC']:.2f}
  BIC       = {stats['BIC']:.2f}
  F-stat    = {stats.get('F_stat', 'N/A')}  (p = {stats.get('F_p', 'N/A')}){hausman_str}{bp_str}{re_str}

Coefficient Table:
{coeff_table}
"""

        sys_prompt = (
            "You are an expert econometrician and data scientist. "
            "Given panel regression output, provide a clear, structured interpretation. "
            "Cover: (1) model choice rationale, (2) coefficient interpretation with economic meaning, "
            "(3) statistical significance and effect sizes, (4) model fit quality, "
            "(5) potential caveats or concerns (endogeneity, heteroskedasticity, etc.), "
            "(6) actionable recommendations. Be concise but insightful. Use plain language."
        )

        col_explain, col_custom = st.columns([3, 2])

        with col_explain:
            if st.button("✦ Generate AI Explanation", type="primary", use_container_width=True):
                with st.spinner("GPT-4 is analysing your results…"):
                    explanation = call_openai(sys_prompt, f"Please explain these panel regression results:\n\n{context}")
                    st.session_state.ai_explanation = explanation

        with col_custom:
            custom_q = st.text_input("Ask a specific question about the results…",
                                      placeholder="e.g. Is x1 economically significant?")
            if st.button("Ask GPT-4", use_container_width=True) and custom_q:
                with st.spinner("Thinking…"):
                    answer = call_openai(
                        sys_prompt,
                        f"Here are the regression results:\n\n{context}\n\nQuestion: {custom_q}",
                    )
                    st.session_state.ai_explanation = answer

        if st.session_state.ai_explanation:
            st.markdown("---")
            st.markdown(f"""
            <div class="ai-label">✦ &nbsp;GPT-4 AI INTERPRETATION</div>
            <div class="ai-box">{st.session_state.ai_explanation}</div>
            """, unsafe_allow_html=True)

        # ── Quick insight cards ────────────────────────────────────────────────
        st.markdown("---")
        st.markdown('<div class="scard-title">Quick Insights</div>', unsafe_allow_html=True)

        ic1, ic2, ic3 = st.columns(3)
        sig_vars = result_df[result_df["p_value"] < 0.05]["Variable"].tolist()
        insig_vars = result_df[result_df["p_value"] >= 0.05]["Variable"].tolist()
        r2_val = stats["R2"]

        with ic1:
            st.markdown(f"""
            <div class="scard">
                <div class="scard-title">Significance</div>
                <div style="color:var(--accent);font-size:1.2rem;font-family:'Syne',sans-serif;font-weight:700;">{len(sig_vars)}/{len(result_df)}</div>
                <div style="color:var(--muted);font-size:0.75rem;margin-top:4px;">variables significant at 5%</div>
                <div style="margin-top:10px;font-size:0.72rem;color:var(--text);">{', '.join(sig_vars) if sig_vars else '—'}</div>
            </div>
            """, unsafe_allow_html=True)
        with ic2:
            r2_color = "#00e5c8" if r2_val > 0.7 else "#f5a623" if r2_val > 0.4 else "#f05c7c"
            r2_label = "Strong fit" if r2_val > 0.7 else "Moderate fit" if r2_val > 0.4 else "Weak fit"
            st.markdown(f"""
            <div class="scard">
                <div class="scard-title">Model Fit</div>
                <div style="color:{r2_color};font-size:1.2rem;font-family:'Syne',sans-serif;font-weight:700;">{r2_val:.4f}</div>
                <div style="color:var(--muted);font-size:0.75rem;margin-top:4px;">R-squared · {r2_label}</div>
            </div>
            """, unsafe_allow_html=True)
        with ic3:
            largest = result_df.iloc[result_df["Coeff"].abs().argmax()]
            st.markdown(f"""
            <div class="scard">
                <div class="scard-title">Largest Effect</div>
                <div style="color:var(--accent2);font-size:1.2rem;font-family:'Syne',sans-serif;font-weight:700;">{largest['Variable']}</div>
                <div style="color:var(--muted);font-size:0.75rem;margin-top:4px;">coeff = {largest['Coeff']:.4f}</div>
            </div>
            """, unsafe_allow_html=True)


# ─── Footer ───────────────────────────────────────────────────────────────────
st.markdown("---")
st.markdown("""
<div style="text-align:center;font-family:'DM Mono',monospace;font-size:0.7rem;color:var(--muted);padding:12px 0;">
    ⬡ PanelStatX · Panel Regression Analysis System · Powered by GPT-4
</div>
""", unsafe_allow_html=True)
